import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_guide_document():
    doc = Document()

    # 페이지 여백 설정 (1인치)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 기본 색상 정의
    PRIMARY_COLOR = RGBColor(15, 76, 129)    # #0F4C81 Navy
    SECONDARY_COLOR = RGBColor(79, 70, 229)  # Indigo
    TEXT_DARK = RGBColor(33, 37, 41)         # Dark Gray
    ACCENT_GREEN = RGBColor(16, 185, 129)    # Emerald

    # 스타일 헬퍼
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(100, 116, 139)
        p.paragraph_format.space_after = Pt(24)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_p(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.bold = True
            run_b.font.size = Pt(10.5)
            run_b.font.color.rgb = TEXT_DARK
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT_DARK
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.2
        return p

    def add_callout(text, title="💡 참고 사항"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        # 배경색 설정 (#F8FAFC)
        shading = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

        # 테두리 설정 (좌측 두꺼운 강조선 #0F4C81)
        borders = parse_xml(r'''
            <w:tcBorders {} >
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="0F4C81"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        '''.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(borders)

        p = cell.paragraphs[0]
        run_t = p.add_run(f"{title}\n")
        run_t.font.bold = True
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = PRIMARY_COLOR

        run_m = p.add_run(text)
        run_m.font.size = Pt(10)
        run_m.font.color.rgb = TEXT_DARK
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(0)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_image_if_exists(img_path, caption="", width_inches=6.2):
        full_path = os.path.abspath(img_path)
        if os.path.exists(full_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            
            run = p.add_run()
            run.add_picture(full_path, width=Inches(width_inches))

            if caption:
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_c = p_cap.add_run(f"▲ {caption}")
                run_c.font.size = Pt(9.5)
                run_c.font.italic = True
                run_c.font.color.rgb = RGBColor(100, 116, 139)
                p_cap.paragraph_format.space_after = Pt(12)
        else:
            add_p(f"[이미지 파일 없음: {img_path}]", space_after=6)

    # --- 1. 문서 제목 및 개요 ---
    add_title("Visit Framework V1.0 종합 사용자 가이드")
    add_subtitle("초보자 5분 구동, 브랜딩 콘솔, SMTP 이메일 설정, 이중 보안 및 1대1 DRM 설치 통합 매뉴얼")

    add_h1("1. 시스템 개요 및 주요 특징")
    add_p("Visit Framework는 기업 및 기관에서 간편하게 도입할 수 있는 차세대 방문자 예약 및 보안 출입 관리 플랫폼입니다. 설치 과정이 자동화되어 있어 누구나 5분 이내에 자신만의 방문자 웹 사이트를 구축할 수 있습니다.")
    
    add_p("• ", "Zero-Config 5분 자동 설치: npm install 명령어 실행 시 환경 변수(.env) 및 SQLite DB가 자동 생성됩니다.")
    add_p("• ", "Dynamic White-Label 브랜딩: 웹 관리자 콘솔(/admin/brand)에서 로고, 회사명, 파비콘을 실시간 변경 가능합니다.")
    add_p("• ", "이중 이메일 보안 인증 (Magic Link): 복잡한 비밀번호 없이 최고 관리자 1회성 매직링크로 보안 접근합니다.")
    add_p("• ", "1대1 Hardware ID DRM: 무단 복제 방지를 위해 서버 하드웨어 MAC 주소와 1대1 매핑하여 저작권을 보호합니다.")

    # --- 2. 5분 초간단 설치 및 PowerShell 해결 ---
    add_h1("2. 5분 초간단 설치 및 구동 순서")
    add_p("컴퓨터 초보자도 쉽게 따라할 수 있도록 아래 4단계 인포그래픽 절차대로 진행하시면 됩니다.")
    
    add_image_if_exists("public/docs/images/visit_ko_installation_flow.png", "그림 1 - 5분 만에 끝나는 초간단 4단계 설치 인포그래픽")

    add_h2("가. 설치 상세 명령어")
    add_p("프로젝트 폴더에서 터미널(Command Prompt 또는 PowerShell)을 열고 아래 명령어를 순서대로 입력합니다:")
    
    add_callout("""# 1. 패키지 다운로드 및 기본 환경(.env) 자동 생성
npm install

# 2. 내 컴퓨터 데이터베이스(SQLite) 자동 생성
npx prisma db push

# 3. 데이터베이스 연결 도구 생성
npx prisma generate

# 4. 내 사이트 켜기!
npm run dev""", "💻 터미널 입력 명령어")

    add_h2("나. 윈도우 PowerShell 보안 권한 오류 해결")
    add_p("윈도우 PowerShell 환경에서 npm 명령어 실행 시 PSSecurityException / UnauthorizedAccess 오류가 발생하는 경우 윈도우 보안 정책에 의한 스크립트 차단 때문입니다.")
    
    add_image_if_exists("public/docs/images/visit_ko_powershell_fix.png", "그림 2 - 윈도우 PowerShell 보안 권한 오류 해결 가이드")

    add_callout("""# PowerShell 권한 해결 명령어 (입력 후 'Y' 엔터)
Set-ExecutionPolicy RemoteSigned -Scope CurrentUser""", "🛠️ PowerShell 해결 명령어")

    # --- 3. 주요 접속 주소 (URL) 안내 ---
    add_h1("3. 주요 접속 주소 (URL) 모음")
    add_p("모든 관리자 콘솔은 무단 URL 접근 및 민감 데이터 유출을 방지하기 위해 이중 보안 인증 게이트웨이로 보호됩니다.")

    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "주요 기능"
    hdr_cells[1].text = "접속 URL"
    hdr_cells[2].text = "설명 및 특징"

    # 테이블 헤더 스타일링
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:fill="0F4C81"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    data = [
        ("🏠 방문객 메인 신청 화면", "http://localhost:3000", "누구나 방문 예약 신청 및 결과를 조회하는 페이지"),
        ("🔒 관리자 통합 보안 로그인", "http://localhost:3000/admin/smtp", "이메일 1-Click 인증으로 세션을 여는 관리자 입구"),
        ("🎨 로고 & 회사명 콘솔", "http://localhost:3000/admin/brand", "내 기업 로고, 회사명, 파비콘, 1대1 라이선스 설정 콘솔")
    ]

    for i, row in enumerate(data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        for c in row_cells:
            c.paragraphs[0].runs[0].font.size = Pt(9.5)
            c.paragraphs[0].runs[0].font.color.rgb = TEXT_DARK
            shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "F8FAFC" if i % 2 == 1 else "FFFFFF"))
            c._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- 4. 방문자 예약 신청 4단계 가이드 ---
    add_h1("4. 방문자 예약 신청 4단계 가이드")
    add_p("방문객은 메인 신청 화면에서 친숙한 마법사(Wizard) 폼을 통해 간편하게 방문을 신청할 수 있습니다.")

    add_image_if_exists("public/docs/images/visit_ko_reservation_wizard.png", "그림 3 - 4단계 방문 예약 마법사 가이드")

    add_p("1. ", "약관 동의: 개인정보 처리방침, 정보보안 서약서 등 5가지 필수 약관에 동의합니다.")
    add_p("2. ", "방문자 정보 입력: 성명, 연락처, 생년월일, 소속 회사명, 차량번호를 입력합니다.")
    add_p("3. ", "접견 담당자 지정: 사내 접견할 담당 임직원의 이름과 이메일 주소를 입력합니다.")
    add_p("4. ", "신청 완료 및 QR 발급: 신청 즉시 6자리 예약번호와 출입용 QR코드가 자동 생성됩니다.")

    # --- 5. 기업 맞춤 브랜딩 콘솔 ---
    add_h1("5. 기업 맞춤 브랜딩 콘솔 가이드")
    add_p("관리자는 웹 화면에서 클릭 몇 번만으로 시스템 내 회사명과 CI 로고, 파비콘을 실시간으로 변경할 수 있습니다.")

    add_image_if_exists("public/docs/images/visit_ko_brand_console.png", "그림 4 - 마스터 브랜딩 및 CI/파비콘 설정 콘솔")

    add_p("• ", "접속 방법: http://localhost:3000/admin/brand 접속 (최고 관리자 메직링크 인증 필요)")
    add_p("• ", "로고 업로드: 투명 배경의 PNG 또는 SVG 로고 파일(최대 2MB)을 드래그 앤 드롭으로 업로드합니다.")
    add_p("• ", "실시간 반영: 설정 저장 시 메인 화면, 푸터, 알림 이메일 템플릿에 실시간 적용됩니다.")

    # --- 6. 이메일(SMTP) 설정 및 자동 알림 ---
    add_h1("6. 이메일(SMTP) 서버 설정 및 5가지 자동 알림 시점")
    add_p("Visit 시스템의 알림 및 로그인 프로세스는 이메일(SMTP)을 기반으로 동작합니다. 설치 후 최초 1회 메일 서버 설정이 필요합니다.")

    add_image_if_exists("public/docs/images/visit_ko_smtp_guide.png", "그림 5 - 이메일(SMTP) 서버 설정 및 5가지 이벤트 알림 트레이")

    add_p("🔔 자동 이메일 알림이 발송되는 5가지 핵심 시점:")
    add_p("1. 방문 신청 접수 시: 방문 신청자에게 예약번호 및 접수 내역 안내 메일 발송")
    add_p("2. 담당자 승인 요청 시: 접견 임직원에게 1-Click 승인/반려 대시보드 보안 링크 발송")
    add_p("3. 최종 승인 완료 시: 방문자 및 동행자 전원에게 출입용 QR코드 및 최종 승인서 안내")
    add_p("4. 방문 예약 취소 시: 방문객 및 담당 임직원에게 취소 통보 메일 발송")
    add_p("5. 차량 정문 입차 시: LPR 카메라인식으로 방문객 차량 통과 시 담당자에게 실시간 알림 메일 발송")

    # --- 7. 이중 보안 게이트웨이 ---
    add_h1("7. 최고 관리자 1-Click 이중 이메일 보안 인증 게이트웨이")
    add_p("외부 사용자가 관리자 URL(/admin/smtp, /admin/brand 등)을 직접 입력하여 접속하더라도, 민감한 설정 정보가 그대로 노출되지 않도록 최고 관리자 이중 이메일 매직링크 인증 레이어가 차단망을 형성합니다.")

    add_image_if_exists("public/docs/images/visit_ko_admin_security_gate.png", "그림 6 - 관리자 1-Click 이중 이메일 보안 인증 게이트웨이")

    add_p("• ", "직접 주소 입력 차단: URL을 직접 입력 시 관리자 인증 폼만 노출되며 민감 정보는 완벽 숨김 처리됩니다.")
    add_p("• ", "1회성 보안 링크: 최고 관리자 메일로 수신된 15분 유효 보안 매직링크를 통해서만 콘솔 조회가 가능합니다.")

    # --- 8. 1대1 Hardware ID DRM 라이선스 해제 가이드 ---
    add_h1("8. 1대1 PC/서버 바인딩 저작권 해제 신청 가이드")
    add_p("본 프로젝트는 MIT License를 따르며, 무단 복제 및 돌려쓰기(재배포)를 방지하기 위해 서버 하드웨어 고유 식별자(Hardware ID: MAC 주소 기반) 1대1 인가 시스템이 적용되어 있습니다.")

    add_image_if_exists("public/docs/images/visit_ko_hw_license_unlock_flow.png", "그림 7 - 1대1 PC/서버 하드웨어 바인딩 저작권 해제 신청 4단계 프로세스")

    add_callout("""1단계: 내 서버 Hardware ID 복사 (http://localhost:3000/admin/brand 접속 하단에서 HW-XXXX-XXXX-XXXX 복사)
2단계: 해제 요청 이메일 발송 (qorwlsgh25@gmail.com 으로 회사명과 Hardware ID 첨부)
3단계: 1대1 전용 해제 코드 수신 (저작권자 검증 후 수신 메일로 VISIT-XXXX-XXXX-XXXX 수신)
4단계: 웹 콘솔 입력 및 저장 (관리자 콘솔 맨 아래 입력란에 입력 후 저장 ➔ ✅ 1대1 정품 인가 완료)""", "🔑 1대1 저작권 해제 4단계 순서")

    add_p("⚠️ ", "무단 복제 방지: 해당 소스코드를 다른 PC나 서버로 통째로 복사할 경우, 타 PC의 Hardware ID가 다르므로 해제 코드가 즉시 무효화되고 저작자 워터마크가 자동으로 다시 강제 노출됩니다.")

    # --- 9. 라이선스 문의 ---
    add_h1("9. 문의 및 라이선스 센터")
    add_p("• ", "저작자 이메일: qorwlsgh25@gmail.com")
    add_p("• ", "오류 제보, 개선 제안, 1대1 라이선스 해제 코드 신청은 위 이메일로 연락주시면 신속하게 지원해 드립니다.")

    output_path = "사용자가이드_V1.0.docx"
    doc.save(output_path)
    print(f"Document created successfully: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    create_guide_document()
